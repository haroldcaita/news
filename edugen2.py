import streamlit as st
import google.generativeai as genai
from docx import Document
from io import BytesIO
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import hashlib
import random
import string
import re
from supabase import create_client, Client
from datetime import datetime

# ---------------------------------------------------------
# Configure Supabase
# ---------------------------------------------------------
NEXT_PUBLIC_SUPABASE_URL = "https://jegrgjhbhhgqlglodheb.supabase.co"
NEXT_PUBLIC_SUPABASE_ANON_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6ImplZ3JnamhiaGhncWxnbG9kaGViIiwicm9sZSI6ImFub24iLCJpYXQiOjE3NDQxOTQ2NzgsImV4cCI6MjA1OTc3MDY3OH0.Cxv3926ZQoKVL0PToITSjc-z9zF767Bzcdyvanpjjg8"
supabase: Client = create_client(NEXT_PUBLIC_SUPABASE_URL, NEXT_PUBLIC_SUPABASE_ANON_KEY)

# ---------------------------------------------------------
# Configure the Google Generative AI API
# ---------------------------------------------------------
genai.configure(api_key="AIzaSyC8zKSK8_xurGABNkBGyn-bbVj4mho-5B8")

generation_config = {
    "temperature": 0.7,
    "top_p": 0.95,
    "top_k": 40,
    "max_output_tokens": 8192,
}

model = genai.GenerativeModel(model_name="gemini-2.0-flash", generation_config=generation_config)

# ---------------------------------------------------------
# Helper Functions
# ---------------------------------------------------------
def send_sms_random(phone_number, verification_code):
    """Simulate sending an SMS with verification code"""
    print(f"Sending verification code {verification_code} to {phone_number}")
    st.success("Verification code sent via SMS!")

def generate_verification_code():
    """Generate a 6-digit verification code"""
    return str(random.randint(100000, 999999))

def hash_password(password):
    """Hash password using SHA-256"""
    return hashlib.sha256(password.encode()).hexdigest()

# ---------------------------------------------------------
# API-based Generation Functions
# ---------------------------------------------------------
def generate_questions(topic, num_questions, qtype):
    if qtype == "Multiple Choice":
        prompt = (f"Generate {num_questions} multiple-choice questions on the topic '{topic}'. "
                  f"Each question should include 4 options (A, B, C, D) and indicate the correct answer.")
    elif qtype == "True or False":
        prompt = (f"Generate {num_questions} true or false questions on the topic '{topic}'. "
                  f"Include the correct answer for each question.")
    elif qtype == "Fill in the Blanks":
        prompt = (f"Generate {num_questions} fill in the blank questions on the topic '{topic}'. "
                  f"Provide the answer for each blank.")
    else:
        prompt = f"Generate {num_questions} questions on the topic '{topic}'."
    try:
        with st.spinner("Generating questions..."):
            response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        st.error(f"Error generating questions: {e}")
        return f"Error generating questions: {e}"

def generate_titles(topic, num_titles):
    prompt = f"Generate {num_titles} research-worthy thesis or capstone project titles for the topic '{topic}'."
    try:
        with st.spinner("Generating titles..."):
            response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        st.error(f"Error generating titles: {e}")
        return f"Error generating titles: {e}"

def generate_reviewer(content, title):
    prompt = f"Generate study materials and key points for the title '{title}' based on the following content: {content}"
    try:
        with st.spinner("Generating reviewer content..."):
            response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        st.error(f"Error generating reviewer content: {e}")
        return f"Error generating reviewer content: {e}"

def generate_essay(essay_input):
    prompt = f"Write an essay based on the following prompt or content:\n\n{essay_input}"
    try:
        with st.spinner("Generating essay..."):
            response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        st.error(f"Error generating essay: {e}")
        return f"Error generating essay: {e}"

def generate_summary(summary_input):
    prompt = f"Provide a clear and concise summary of the following text:\n\n{summary_input}"
    try:
        with st.spinner("Generating summary..."):
            response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        st.error(f"Error generating summary: {e}")
        return f"Error generating summary: {e}"

def generate_citations(style, source):
    prompt = f"Generate a citation for the source '{source}' in {style} format."
    try:
        with st.spinner("Generating citation..."):
            response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        st.error(f"Error generating citation: {e}")
        return f"Error generating citation: {e}"

# ---------------------------------------------------------
# Export Functions
# ---------------------------------------------------------
def export_docx(content):
    doc = Document()
    doc.add_heading("Generated Content", level=0)
    doc.add_paragraph(content)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def export_pdf(content):
    buffer = BytesIO()
    can = canvas.Canvas(buffer, pagesize=letter)
    text_object = can.beginText(40, 750)
    for line in content.split('\n'):
        text_object.textLine(line)
    can.drawText(text_object)
    can.showPage()
    can.save()
    buffer.seek(0)
    return buffer

# ---------------------------------------------------------
# Authentication Functions
# ---------------------------------------------------------
def create_user(username, password, contact=None):
    hashed_password = hash_password(password)
    try:
        data = supabase.table("users").insert({
            "username": username,
            "password": hashed_password,
            "contact": contact,
            "created_at": str(datetime.now())
        }).execute()
        return True
    except Exception as e:
        st.error(f"Error creating user: {e}")
        return False

def authenticate_user(username, password):
    hashed_password = hash_password(password)
    try:
        data = supabase.table("users").select("*").eq("username", username).eq("password", hashed_password).execute()
        return len(data.data) > 0
    except Exception as e:
        st.error(f"Error authenticating user: {e}")
        return False

def forgot_password_flow():
    """Handle the password reset flow with SMS verification"""
    st.subheader("Forgot Password")

    # Initialize session state variables
    if 'verification_code' not in st.session_state:
        st.session_state.verification_code = None
    if 'reset_username' not in st.session_state:
        st.session_state.reset_username = None
    if 'show_verification_input' not in st.session_state:
        st.session_state.show_verification_input = False
    if 'show_password_reset' not in st.session_state:
        st.session_state.show_password_reset = False

    # Step 1: Username input
    username = st.text_input("Enter your username", key="forgot_username")

    if st.button("Send Verification Code", key="send_verification"):
        # Check if username exists in Supabase
        try:
            response = supabase.table("users").select("username, contact").eq("username", username).execute()

            if response.data and len(response.data) > 0:
                user = response.data[0]

                # Check if user has contact information
                if user.get('contact'):
                    # Generate and store verification code
                    verification_code = generate_verification_code()
                    st.session_state.verification_code = verification_code
                    st.session_state.reset_username = username

                    # Simulate sending SMS
                    send_sms_random(user['contact'], verification_code)

                    # Show verification input section
                    st.session_state.show_verification_input = True
                    st.success("Verification code sent to your registered phone number")
                else:
                    st.error("No contact information found for this user. Please contact support.")
            else:
                st.error("Username not found. Please try again.")
        except Exception as e:
            st.error(f"Error looking up user: {e}")

    # Step 2: Verification code input
    if st.session_state.show_verification_input:
        verification_input = st.text_input("Enter verification code", key="verification_input")

        if st.button("Verify Code", key="verify_code"):
            if verification_input == st.session_state.verification_code:
                st.success("Verification successful!")
                st.session_state.show_password_reset = True
            else:
                st.error("Invalid verification code. Please try again.")

    # Step 3: Password reset
    if st.session_state.show_password_reset:
        new_password = st.text_input("New password", type="password", key="new_password")
        confirm_password = st.text_input("Confirm new password", type="password", key="confirm_password")

        if st.button("Reset Password", key="reset_password"):
            if new_password == confirm_password:
                try:
                    # Update password in Supabase
                    hashed_password = hash_password(new_password)
                    response = supabase.table("users").update(
                        {"password": hashed_password}
                    ).eq("username", st.session_state.reset_username).execute()

                    if response.data:
                        st.success("Password reset successfully! You can now login with your new password.")

                        # Clear session state
                        st.session_state.verification_code = None
                        st.session_state.reset_username = None
                        st.session_state.show_verification_input = False
                        st.session_state.show_password_reset = False

                        # Rerun to clear the form
                        st.rerun()
                    else:
                        st.error("Password reset failed. Please try again.")
                except Exception as e:
                    st.error(f"Error resetting password: {e}")
            else:
                st.error("Passwords don't match. Please try again.")

    # Back to login link
    st.markdown("[Back to login](#login)")

# ---------------------------------------------------------
# Initialize Session State
# ---------------------------------------------------------
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
if "username" not in st.session_state:
    st.session_state.username = ""
if "generated_content" not in st.session_state:
    st.session_state.generated_content = ""
if "history" not in st.session_state:
    st.session_state.history = []
if "recent_generated" not in st.session_state:
    st.session_state.recent_generated = []

# ---------------------------------------------------------
# App Title & Tagline
# ---------------------------------------------------------
st.set_page_config(page_title="EduGen: Smart Academic Generator", page_icon="🎓")
st.title("EduGen: Smart Academic Generator")
st.markdown("*Making Learning & Teaching Easier* :mortar_board:")
st.markdown("---")

# ---------------------------------------------------------
# User Login/Signup Section
# ---------------------------------------------------------
if not st.session_state.logged_in:
    with st.container():
        st.header("User Login/Signup")
        login_tab, signup_tab, forgot_tab = st.tabs(["Login", "Signup", "Forgot Password"])

        with login_tab:
            username_login = st.text_input("Username", key="username_login")
            password_login = st.text_input("Password", type="password", key="password_login")
            if st.button("Login", key="login_button"):
                if authenticate_user(username_login, password_login):
                    st.session_state.logged_in = True
                    st.session_state.username = username_login
                    st.success("Logged in successfully!", icon="✅")
                    st.rerun()
                else:
                    st.error("Invalid credentials.", icon="🚨")

        with signup_tab:
            username_signup = st.text_input("Username", key="username_signup")
            password_signup = st.text_input("Password", type="password", key="password_signup")
            contact_signup = st.text_input("Phone Number (for password recovery)", key="contact_signup")
            if st.button("Signup", key="signup_button"):
                if create_user(username_signup, password_signup, contact_signup):
                    st.success("Account created! You can now log in.", icon="✅")
                else:
                    st.error("Error creating account. Username may already exist.", icon="🚨")

        with forgot_tab:
            forgot_password_flow()
    st.stop()

# ---------------------------------------------------------
# Main App Content after Login
# ---------------------------------------------------------
tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs(["Question Generator", "Title Generator", "Reviewer Creator", "Essay Generator", "Summary Generator", "Citation & Bibliography Tool"])

def save_generated_content(username, content, generation_type, topic=None):
    """Saves the generated content to the database."""
    try:
        supabase.table("generated_content").insert({
            "username": username,
            "content": content,
            "created_at": str(datetime.now()),
            "generation_type": generation_type,
            "topic": topic,
        }).execute()
    except Exception as e:
        st.error(f"Error saving generated content: {e}")

with tab1:
    st.header("Question Generator")
    topic = st.text_input("Enter the topic for questions", placeholder="e.g., Photosynthesis")
    num_questions = st.number_input("Number of Questions", min_value=1, max_value=20, value=5)
    qtype = st.selectbox("Select Question Type", ("Multiple Choice", "True or False", "Fill in the Blanks"))

    if st.button("Generate Questions"):
        content = generate_questions(topic, num_questions, qtype)

        st.session_state.generated_content = content
        st.session_state.history.append(content)
        st.session_state.recent_generated.append({"type": "Question", "content": content[:50] + "...", "full_content": content})
        save_generated_content(st.session_state.username, content, "Question Generator", topic)
        st.rerun() # Rerun to display the content below

with tab2:
    st.header("Thesis & Capstone Title Generator")
    topic = st.text_input("Enter the topic for titles", placeholder="e.g., Impact of Social Media on Teenagers")
    num_titles = st.number_input("Number of Titles", min_value=1, max_value=10, value=3)
    if st.button("Generate Titles"):
        content = generate_titles(topic, num_titles)
        st.session_state.generated_content = content
        st.session_state.history.append(content)
        st.session_state.recent_generated.append({"type": "Title", "content": content[:50] + "...", "full_content": content})
        save_generated_content(st.session_state.username, content, "Title Generator", topic)
        st.rerun()

with tab3:
    st.header("Reviewer Creator")
    title_input = st.text_input("Enter the title for the study material", placeholder="e.g., Chapter 3 - Cell Biology")
    content_input = st.text_area("Enter content or topics", height=150, placeholder="Key concepts, formulas, etc.")
    if st.button("Generate Reviewer"):
        content = generate_reviewer(content_input, title_input)
        st.session_state.generated_content = content
        st.session_state.history.append(content)
        st.session_state.recent_generated.append({"type": "Reviewer", "content": content[:50] + "...", "full_content": content, "topic": title_input})
        save_generated_content(st.session_state.username, content, "Reviewer Creator", title_input)
        st.rerun()

with tab4:
    st.header("Essay Generator")
    essay_input = st.text_area("Enter your essay prompt or content", height=250, placeholder="Write your essay prompt or paste your essay content here...")
    if st.button("Generate Essay"):
        content = generate_essay(essay_input)
        st.session_state.generated_content = content
        st.session_state.history.append(content)
        st.session_state.recent_generated.append({"type": "Essay", "content": content[:50] + "...", "full_content": content, "topic": essay_input[:50] + "..."})
        save_generated_content(st.session_state.username, content, "Essay Generator", essay_input[:100])
        st.rerun()

with tab5:
    st.header("Summary Generator")
    summary_input = st.text_area("Enter text to summarize", height=250, placeholder="Paste the text you want to summarize here...")
    if st.button("Generate Summary"):
        content = generate_summary(summary_input)
        st.session_state.generated_content = content
        st.session_state.history.append(content)
        st.session_state.recent_generated.append({"type": "Summary", "content": content[:50] + "...", "full_content": content, "topic": summary_input[:50] + "..."})
        save_generated_content(st.session_state.username, content, "Summary Generator", summary_input[:100])
        st.rerun()

with tab6:
    st.header("Citation & Bibliography Tool")
    source = st.text_input("Enter the source title or URL", placeholder="e.g., www.example.com or 'The Great Gatsby'")
    style = st.selectbox("Select Citation Style", ("APA", "MLA", "Chicago"))
    if st.button("Generate Citation"):
        content = generate_citations(style, source)
        st.session_state.generated_content = content
        st.session_state.history.append(content)
        st.session_state.recent_generated.append({"type": "Citation", "content": content, "full_content": content, "topic": source, "style": style})
        save_generated_content(st.session_state.username, content, "Citation Generator", source + f" ({style})")
        st.rerun()

# Sidebar: User Info and Logout
with st.sidebar:
    st.image("https://cdn-icons-png.flaticon.com/512/149/149071.png", width=50)
    st.write(f"Logged in as: {st.session_state.username}")
    if st.button("Logout"):
        st.session_state.logged_in = False
        st.session_state.username = ""
        st.session_state.generated_content = ""
        st.session_state.history = []
        st.session_state.recent_generated = []
        st.rerun()

    st.markdown("---")

    if st.checkbox("Show Recent Generated"):
        st.header("Recent Generated")
        try:
            response = supabase.table("generated_content").select("id, generation_type, content, created_at").eq("username", st.session_state.username).order("created_at", desc=True).limit(5).execute()
            recent_items = response.data
            if recent_items:
                recent_list = []
                recent_full_content = {}
                for item in recent_items:
                    title = f"{item.get('generation_type', 'Content')} (Created: {item['created_at'][:16]})"
                    preview = item.get('content', 'No content')[:80] + ("..." if len(item.get('content', '')) > 80 else "")
                    recent_list.append(preview)
                    recent_full_content[preview] = item.get('content', 'No content')

                if recent_list:
                    selected_preview = st.selectbox("View Recent Generations", recent_list)
                    st.subheader("Full Content:")
                    st.write(recent_full_content.get(selected_preview, "No content selected."))
                else:
                    st.info("No recent generated content.")
            else:
                st.info("No recent generated content.")
        except Exception as e:
            st.error(f"Error fetching recent generations: {e}")


# ---------------------------------------------------------
# Display Generated Content and Edit Area
# ---------------------------------------------------------
if st.session_state.generated_content:
    st.header("Generated Content")
    content = st.text_area("Generated Output", value=st.session_state.generated_content, height=300)
    st.session_state.generated_content = content

# ---------------------------------------------------------
# Export Options: Download as DOCX or PDF
# ---------------------------------------------------------
if st.session_state.generated_content:
    st.header("Export Options")
    col1, col2 = st.columns(2)
    with col1:
        docx_buffer = export_docx(st.session_state.generated_content)
        st.download_button(
            label="Download as DOCX",
            data=docx_buffer,
            file_name="generated_content.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
    with col2:
        pdf_buffer = export_pdf(st.session_state.generated_content)
        st.download_button(
            label="Download as PDF",
            data=pdf_buffer,
            file_name="generated_content.pdf",
            mime="application/pdf",
            use_container_width=True
        )
    st.markdown("---")